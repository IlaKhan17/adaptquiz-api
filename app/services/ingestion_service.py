import asyncio
import io
import re
from uuid import uuid4

from fastapi import HTTPException, UploadFile
from pypdf import PdfReader
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.chunker import chunk_text
from app.core.vector_store import add_texts_and_save
from app.models.document import Document
from app.schemas.ingest import IngestResponse

_MIN_TEXT_LENGTH = 100

# Patterns that indicate non-academic content (TOC, index, page numbers)
_TOC_LINE = re.compile(r'.+\.{3,}\s*\d+\s*$')           # "Chapter 3 ........ 45"
_PAGE_NUMBER = re.compile(r'^\s*(page\s*)?\d+\s*$', re.IGNORECASE)  # "12" or "Page 12"
_PAGE_X_OF_Y = re.compile(r'^\s*\d+\s+of\s+\d+\s*$', re.IGNORECASE)  # "3 of 10"
_INDEX_ENTRY = re.compile(r'^[A-Za-z][A-Za-z\s\-,]{1,50}[\d,\s\-–]+$')  # "photosynthesis, 45, 67"
_SECTION_STUB = re.compile(r'^\s*(\d+\.)+\s*$')          # "3.1.2." with no text


def _clean_extracted_text(text: str) -> str:
    """Strip TOC lines, standalone page numbers, and index entries from extracted text."""
    cleaned = []
    for line in text.split('\n'):
        s = line.strip()
        if not s:
            cleaned.append('')
            continue
        if _PAGE_NUMBER.match(s):
            continue
        if _PAGE_X_OF_Y.match(s):
            continue
        if _TOC_LINE.match(s):
            continue
        if _SECTION_STUB.match(s):
            continue
        # Index entries: very short lines that are mostly numbers and commas
        if _INDEX_ENTRY.match(s) and sum(c.isdigit() or c == ',' for c in s) / max(len(s), 1) > 0.35:
            continue
        cleaned.append(line)

    result = '\n'.join(cleaned)
    # Collapse runs of 3+ blank lines into 2
    return re.sub(r'\n{3,}', '\n\n', result).strip()


async def ingest_document(
    file: UploadFile,
    subject: str,
    user_id: str,
    db: AsyncSession,
) -> IngestResponse:
    raw_bytes = await file.read()
    filename = file.filename or "upload"

    text = _clean_extracted_text(_extract_text(raw_bytes, filename))

    if len(text) < _MIN_TEXT_LENGTH:
        raise HTTPException(
            status_code=422,
            detail=f"Extracted text is too short ({len(text)} chars). "
                   "Minimum is 100 characters — check the file has readable content.",
        )

    doc_id = str(uuid4())
    chunks = chunk_text(text, doc_id)

    texts = [c["text"] for c in chunks]
    metadatas = [
        {**c["metadata"], "subject": subject, "filename": filename}
        for c in chunks
    ]

    await asyncio.to_thread(add_texts_and_save, texts, metadatas)

    document = Document(
        user_id=user_id,
        doc_id=doc_id,
        filename=filename,
        subject=subject,
        chunks_created=len(chunks),
        total_chars=len(text),
    )
    db.add(document)
    await db.commit()

    return IngestResponse(
        doc_id=doc_id,
        filename=filename,
        subject=subject,
        chunks_created=len(chunks),
        total_chars=len(text),
    )


def _extract_text(raw_bytes: bytes, filename: str) -> str:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        try:
            reader = PdfReader(io.BytesIO(raw_bytes))
            return "\n".join(page.extract_text() or "" for page in reader.pages).strip()
        except Exception as exc:
            raise HTTPException(
                status_code=422,
                detail=f"Could not read PDF: {exc}",
            ) from exc
    if lower.endswith(".txt"):
        try:
            return raw_bytes.decode("utf-8").strip()
        except UnicodeDecodeError as exc:
            raise HTTPException(
                status_code=422,
                detail="Text file is not valid UTF-8.",
            ) from exc
    if lower.endswith(".docx"):
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(io.BytesIO(raw_bytes))
            return "\n".join(
                para.text for para in doc.paragraphs if para.text.strip()
            ).strip()
        except Exception as exc:
            raise HTTPException(
                status_code=422,
                detail=f"Could not read DOCX: {exc}",
            ) from exc
    raise HTTPException(
        status_code=422,
        detail=f"Unsupported file type '{filename}'. Accepted formats: PDF, DOCX, TXT.",
    )
